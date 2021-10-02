import random

from docx import Document
from docx.shared import Inches

from csv import DictReader

from config import version

document = Document()

# displays info about program
print(f"RLD {version}")
print("Written by Florencio\n")

# sets document title
document.add_heading('Generic Test Title', 0)

# sets paragraph under title
document.add_paragraph('Generic description for said test ')

# set current question number to one
i = 1

# asks for file name
input_file = input("Whats the name of the file you want to convert?\n")

output_file = input("Whats the name you'd like to name the document?\n")

# opens test.csv as read_obj
with open(f"{input_file}.csv", 'r') as read_obj:
		csv_dict_reader = DictReader(read_obj)

		# begins to read each row
		for row in csv_dict_reader:

			# adds heading, which is the question
			document.add_heading(f"{i}. {row['Question']}", level=1)

			# lets the user know what question its working on
			print(f"Working on question {i}..")

			# add all the questions into a list
			answers = list([row['Correct'], row['Distractor1'], row['Distractor2'], row['Distractor3']])

			# shuffles questions in list
			random.seed()
			random.shuffle(answers)

			# adds each answer as a paragraph
			document.add_paragraph(f"A. {answers[0]}", style='List')
			document.add_paragraph(f"B. {answers[1]}", style='List')
			document.add_paragraph(f"C. {answers[2]}", style='List')
			document.add_paragraph(f"D. {answers[3]}", style='List')

			# lets the user know that the question has been completed
			print(f"Question {i} is complete!")

			# sets current question number with an additional one
			i = i + 1


section = document.sections[0]
 
# adds footer
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "\tGenerated with love by Florencio"

# save document
document.save(f"{output_file}.docx")

# lets the user know that the document is now complete
print(f"Document is now complete! Saved as '{output_file}.docx'\n")
input("Press enter to end the program")